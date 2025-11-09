# src/utils/markdown_converter.py

import io
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from markdown_it import MarkdownIt
from xhtml2pdf import pisa

def add_run_with_inline_formatting(paragraph: Paragraph, text: str, bold: bool = False, italic: bool = False, code: bool = False):
    """Adiciona um 'run' a um parágrafo, aplicando formatação inline."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = 'Courier New'

def convert_markdown_to_docx_stream(md_text: str) -> io.BytesIO:
    """
    Converte uma string Markdown para um stream de bytes de um arquivo DOCX,
    suportando títulos, parágrafos, listas, formatação inline, tabelas,
    linhas horizontais e blocos de código.
    """
    document = Document()
    md = MarkdownIt()
    tokens = md.parse(md_text)

    current_paragraph = None
    list_level = 0
    table_data = []
    in_table = False

    for i, token in enumerate(tokens):
        # --- Títulos ---
        if token.type == 'heading_open':
            level = int(token.tag[1])
            title_text = tokens[i+1].content.strip()
            document.add_heading(title_text, level=min(level, 4))
            current_paragraph = None
        
        # --- Parágrafos Normais ---
        elif token.type == 'paragraph_open':
            # Parágrafos dentro de tabelas ou listas são tratados de forma diferente
            if not in_table:
                current_paragraph = document.add_paragraph()
                current_paragraph.paragraph_format.space_after = Pt(6)

        # --- Listas (Lógica existente) ---
        elif token.type in ['bullet_list_open', 'ordered_list_open']:
            list_level += 1
        elif token.type in ['bullet_list_close', 'ordered_list_close']:
            list_level = max(0, list_level - 1)

        # --- Linha Horizontal ---
        elif token.type == 'hr':
            p = document.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement('w:pBdr')
            p_pr.append(p_bdr)
            bottom_bdr = OxmlElement('w:bottom')
            bottom_bdr.set(qn('w:val'), 'single')
            bottom_bdr.set(qn('w:sz'), '6')
            bottom_bdr.set(qn('w:space'), '1')
            bottom_bdr.set(qn('w:color'), 'auto')
            p_bdr.append(bottom_bdr)
            current_paragraph = None

        # --- Blocos de Código ---
        elif token.type == 'fence': # Bloco de código com ```
            p = document.add_paragraph(style='No Spacing')
            run = p.add_run(token.content.strip('\n'))
            run.font.name = 'Courier New'
            current_paragraph = None
            
        # --- Tabelas ---
        elif token.type == 'table_open':
            in_table = True
            table_data = []
        elif token.type == 'tr_open':
            table_data.append([]) # Adiciona uma nova linha
        elif token.type == 'th_open' or token.type == 'td_open':
            # O conteúdo da célula está no próximo token 'inline'
            cell_content = tokens[i+1].content
            table_data[-1].append(cell_content) # Adiciona célula à linha atual
        elif token.type == 'table_close':
            in_table = False
            if table_data:
                # Cria a tabela no documento do Word
                num_rows = len(table_data)
                num_cols = len(table_data) if num_rows > 0 else 0
                table = document.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                
                # Preenche as células
                for row_idx, row_content in enumerate(table_data):
                    for col_idx, cell_content in enumerate(row_content):
                        cell = table.cell(row_idx, col_idx)
                        # Limpa o parágrafo padrão e adiciona o conteúdo
                        cell.text = cell_content
                
                document.add_paragraph() # Adiciona um espaço após a tabela
            table_data = []
            current_paragraph = None
            
        # --- Conteúdo (Texto e Formatação Inline) ---
        elif token.type == 'inline' and (token.content or token.children):
            # Verificamos se o "avô" deste token é um item de lista
            is_in_list = (i > 1) and tokens[i-2].type == 'list_item_open'
            
            if is_in_list:
                list_type_token = tokens[i-3]
                style = 'List Number' if list_type_token.type == 'ordered_list_open' else 'List Bullet'
                p = document.add_paragraph(style=style)
                
                is_bold, is_italic, is_code = False, False, False
                for child in token.children:
                    if child.type == 'strong_open': is_bold = True
                    elif child.type == 'strong_close': is_bold = False
                    elif child.type == 'em_open': is_italic = True
                    elif child.type == 'em_close': is_italic = False
                    elif child.type == 'code_inline':
                        add_run_with_inline_formatting(p, child.content, code=True)
                    elif child.type == 'text':
                        add_run_with_inline_formatting(p, child.content, bold=is_bold, italic=is_italic)
                
                p.paragraph_format.space_after = Pt(0)
                if list_level > 1:
                    p.paragraph_format.left_indent = Inches(0.25 * (list_level - 1))
            
            # Se for um parágrafo normal (e não dentro de uma tabela)
            elif current_paragraph is not None and not in_table:
                is_bold, is_italic, is_code = False, False, False
                for child in token.children:
                    if child.type == 'strong_open': is_bold = True
                    elif child.type == 'strong_close': is_bold = False
                    elif child.type == 'em_open': is_italic = True
                    elif child.type == 'em_close': is_italic = False
                    elif child.type == 'code_inline':
                        add_run_with_inline_formatting(current_paragraph, child.content, code=True)
                    elif child.type == 'text':
                        add_run_with_inline_formatting(current_paragraph, child.content, bold=is_bold, italic=is_italic)
                    elif child.type == 'softbreak':
                        current_paragraph.add_run('\n')

    stream = io.BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream

def convert_markdown_to_pdf_stream(md_text: str) -> io.BytesIO:
    """Converte uma string Markdown para um stream de bytes de um arquivo PDF."""
    # Converte o Markdown para HTML
    md = MarkdownIt()
    html_text = md.render(md_text)
    
    # Prepara o stream de bytes para o resultado
    stream = io.BytesIO()
    
    # Usa a biblioteca pisa (parte do xhtml2pdf) para converter HTML para PDF
    pisa_status = pisa.CreatePDF(
        io.StringIO(html_text),  # Fonte do HTML
        dest=stream             # Destino (o stream de bytes)
    )
    
    if pisa_status.err:
        raise Exception("Erro ao converter HTML para PDF")
        
    stream.seek(0)
    return stream