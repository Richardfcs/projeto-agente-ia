# /src/tasks/file_generators.py
import io
import logging
from typing import Iterable, List, Optional, Any

from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

logger = logging.getLogger(__name__)


# NOTA: Alterei o tipo para Iterable[Any] para ser um pouco mais explícito,
# já que a função lida com a conversão de qualquer tipo para string.
def _normalize_topicos(topicos: Optional[Iterable[Any]]) -> List[str]:
    """Normaliza entrada: aceita None, iteráveis; converte cada item para str e filtra None."""
    if topicos is None:
        return []
    try:
        # A lógica aqui já é excelente, convertendo qualquer coisa para string.
        return ["" if t is None else str(t) for t in topicos]
    except TypeError:
        # A mensagem de erro é clara e útil. Ótimo!
        raise ValueError("topicos deve ser um iterável ou None")


def criar_docx_stream(
    topicos: Optional[Iterable[Any]] = None,
    filename: str = "relatorio.docx",
    title: str = "Relatório Gerado por IA"  # SUGESTÃO: Adicionado parâmetro de título.
) -> io.BytesIO:
    """
    Cria DOCX em memória e retorna BytesIO.
    topicos: iterável de strings.
    filename: opcional (atribui stream.name).
    title: título principal do documento.
    """
    topicos = _normalize_topicos(topicos)
    stream = io.BytesIO()
    try:
        doc = Document()
        doc.add_heading(title, 0)  # Usa o título customizável.
        for topico in topicos:
            # NOTA: Para conteúdo mais rico, aqui seria o lugar para adicionar lógica
            # para, por exemplo, criar parágrafos, listas com marcadores, etc.
            doc.add_paragraph(topico)
        doc.save(stream)
        stream.seek(0)
        stream.name = filename
        return stream
    except Exception as e:
        logger.exception("Erro ao criar DOCX")
        raise RuntimeError(f"Falha ao gerar DOCX: {e}") from e


def criar_xlsx_stream(
    topicos: Optional[Iterable[Any]] = None,
    filename: str = "relatorio.xlsx",
    title: str = "Relatório IA",
    write_only: bool = False  # Mantendo o parâmetro para consistência
) -> io.BytesIO:
    """
    Cria XLSX em memória e retorna BytesIO.
    Interpreta cada 'topico' como uma linha e separa as colunas pelo caractere ';'.
    """
    topicos = _normalize_topicos(topicos)
    stream = io.BytesIO()
    try:
        wb = Workbook(write_only=write_only)
        ws = wb.create_sheet(title=title) if write_only else wb.active
        if not write_only:
            ws.title = title

        for linha_texto in topicos:
            # --- INÍCIO DA MUDANÇA ---
            # Adiciona a linha apenas se ela tiver conteúdo real
            if linha_texto and linha_texto.strip():
                # Divide a linha de texto pelo separador para criar as colunas
                colunas = [celula.strip() for celula in linha_texto.split(';')]
                # O método .append() do openpyxl aceita uma lista e a distribui pelas colunas
                ws.append(colunas)
            # --- FIM DA MUDANÇA ---

        wb.save(stream)
        stream.seek(0)
        stream.name = filename
        return stream
    except Exception as e:
        logger.exception("Erro ao criar XLSX")
        raise RuntimeError(f"Falha ao gerar XLSX: {e}") from e


def criar_pdf_stream(
    topicos: Optional[Iterable[Any]] = None,
    filename: str = "relatorio.pdf",
    title: str = "Relatório Gerado por IA"
) -> io.BytesIO:
    """
    Cria PDF em memória e retorna BytesIO.
    Para templates PDF, sugiro usar outra função que preencha formulários (pdfrw/pypdf).
    """
    topicos = _normalize_topicos(topicos)
    stream = io.BytesIO()
    try:
        doc = SimpleDocTemplate(stream, pagesize=letter)
        styles = getSampleStyleSheet()
        # A construção dos "flowables" está perfeita para reportlab.
        flowables = [Paragraph(title, styles['h1']), Spacer(1, 12)]
        for topico in topicos:
            flowables.append(Paragraph(topico, styles['Normal']))
            flowables.append(Spacer(1, 6))
        doc.build(flowables)
        stream.seek(0)
        stream.name = filename
        return stream
    except Exception as e:
        logger.exception("Erro ao criar PDF")
        raise RuntimeError(f"Falha ao gerar PDF: {e}") from e