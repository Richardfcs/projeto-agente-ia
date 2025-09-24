# Arquivo: /src/tasks/ia_processor.py

import google.generativeai as genai
from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import io
from src.config import Config
from src.db.mongo import get_db, get_gridfs
from bson import ObjectId
from datetime import datetime

# Configuração da API do Gemini
if not Config.GOOGLE_API_KEY:
    raise ValueError("A variável de ambiente GOOGLE_API_KEY não foi definida.")
genai.configure(api_key=Config.GOOGLE_API_KEY)


def gerar_resposta(historico_mensagens: list) -> str:
    """
    Interage com o Gemini usando um histórico de conversa completo e retorna a nova resposta.
    """
    modelo = genai.GenerativeModel("gemini-2.5-flash-lite")
    
    # O Gemini espera um formato específico para o histórico.
    # O papel do assistente é 'model'. O papel do usuário é 'user'.
    mensagens_para_api = []
    for msg in historico_mensagens:
        role = "model" if msg["role"] == "assistant" else "user"
        mensagens_para_api.append({"role": role, "parts": [msg["content"]]})

    # Inicia o chat com o histórico
    chat = modelo.start_chat(history=mensagens_para_api[:-1]) # Envia todo o histórico, exceto a última mensagem
    
    # Envia a última mensagem do usuário para obter a nova resposta
    ultima_mensagem = mensagens_para_api[-1]["parts"]
    resposta = chat.send_message(ultima_mensagem)
    
    return resposta.text

def extrair_topicos(texto: str) -> list:
    """Extrai os tópicos da resposta para formatação."""
    return [linha.strip() for linha in texto.split("\n") if linha.strip()]

def criar_docx_stream(topicos: list) -> io.BytesIO:
    """Cria um arquivo DOCX em memória e retorna seu stream de bytes."""
    doc = Document()
    doc.add_heading("Relatório Gerado por IA", 0)
    for topico in topicos:
        doc.add_paragraph(topico)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

def criar_xlsx_stream(topicos: list) -> io.BytesIO:
    """Cria um arquivo XLSX em memória e retorna seu stream de bytes."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Relatório IA"
    for i, topico in enumerate(topicos, start=1):
        ws[f"A{i}"] = topico
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream

def criar_pdf_stream(topicos: list) -> io.BytesIO:
    """Cria um arquivo PDF em memória e retorna seu stream de bytes."""
    stream = io.BytesIO()
    doc = SimpleDocTemplate(stream, pagesize=letter)
    styles = getSampleStyleSheet()
    flowables = [Paragraph("Relatório Gerado por IA", styles['Heading1']), Spacer(1, 12)]
    
    for topico in topicos:
        flowables.append(Paragraph(topico, styles['Normal']))
        flowables.append(Spacer(1, 6))
        
    doc.build(flowables)
    stream.seek(0)
    return stream


# Esta não é mais uma tarefa do Celery, é uma função Python padrão.
def processar_solicitacao_ia(message_id: str) -> str:
    """
    Executa o fluxo completo de processamento de IA, agora com memória de conversa.
    """
    print(f"Iniciando processamento com memória para a mensagem: {message_id}")
    db = get_db()
    fs = get_gridfs()
    
    try:
        # 1. Buscar a mensagem atual para obter o ID da conversa
        mensagem_atual = db.messages.find_one({"_id": ObjectId(message_id)})
        if not mensagem_atual:
            print(f"ERRO: Mensagem com ID {message_id} não encontrada.")
            return "Falha"

        conversation_id = mensagem_atual.get("conversation_id")

        # 2. Buscar TODO o histórico de mensagens da conversa, em ordem
        historico_cursor = db.messages.find(
            {"conversation_id": conversation_id}
        ).sort("timestamp", 1)
        
        # Converte o cursor em uma lista de dicionários
        historico_completo = list(historico_cursor)
        
        # 3. Chamar a IA com o histórico completo
        resposta_bruta = gerar_resposta(historico_completo)
        topicos = extrair_topicos(resposta_bruta)
        
        # O resto do fluxo permanece o mesmo...
        # 4. Criar o arquivo em memória
        nome_arquivo = "relatorio_gerado_com_contexto.docx"
        arquivo_stream = criar_docx_stream(topicos)
        
        # 5. Salvar o arquivo no GridFS
        output_file_id = fs.put(arquivo_stream, filename=nome_arquivo)
        print(f"Arquivo com contexto salvo no GridFS com ID: {output_file_id}")
        
        # 6. Criar o metadado do arquivo
        output_doc_meta = {
            "filename": nome_arquivo,
            "gridfs_file_id": output_file_id,
            "owner_id": mensagem_atual.get("user_id"),
            "created_at": datetime.utcnow()
        }
        output_doc = db.documents.insert_one(output_doc_meta)
        
        # 7. Criar a mensagem de resposta do assistente
        assistant_message = {
            "conversation_id": conversation_id,
            "role": "assistant",
            "content": f"Com base em nossa conversa, gerei o documento '{nome_arquivo}'.",
            "generated_document_id": output_doc.inserted_id,
            "user_id": mensagem_atual.get("user_id"),
            "timestamp": datetime.utcnow()
        }
        db.messages.insert_one(assistant_message)
        
        # 8. Atualizar a conversa
        db.conversations.update_one(
            {"_id": conversation_id},
            {"$set": {"last_updated_at": datetime.utcnow()}}
        )

        print(f"Processamento com memória para a mensagem {message_id} concluído.")
        return "Sucesso"

    except Exception as e:
        print(f"ERRO CRÍTICO ao processar a mensagem {message_id}: {e}")
        # Salva uma mensagem de erro no banco de dados para o usuário ver.
        db.messages.insert_one({
            "conversation_id": message.get("conversation_id"),
            "role": "assistant",
            "content": f"Ocorreu um erro ao processar sua solicitação. Por favor, tente novamente.",
            "user_id": message.get("user_id"),
            "timestamp": datetime.utcnow(),
            "is_error": True
        })
        return "Falha"